#!/usr/bin/env python3
"""
Cortex Example: docx-report-from-json
Description: Generate styled Word reports from JSON structure
Tags: word,docx,report,template
Captured: 2026-03-25
Source: docx-report-from-template.py

Usage:
  python ~/.claude/skills/cortex/cookbook/docx-report-from-json.py
"""

#!/usr/bin/env python3
"""
Generate a Word document report with sections, tables, and styled headings.
Accepts a JSON structure defining sections.

Usage:
    python docx-report-from-template.py report_data.json [output.docx]

JSON format:
{
    "title": "Monthly Report",
    "subtitle": "March 2026",
    "author": "Team Name",
    "sections": [
        {
            "heading": "Summary",
            "content": "This month we achieved..."
        },
        {
            "heading": "Metrics",
            "table": {
                "headers": ["Metric", "Value", "Change"],
                "rows": [
                    ["Revenue", "$150K", "+12%"],
                    ["Users", "5,200", "+8%"]
                ]
            }
        }
    ]
}
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_report(data, output_path):
    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # Title
    title = doc.add_heading(data["title"], level=0)
    title.runs[0].font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    if data.get("subtitle"):
        p = doc.add_paragraph(data["subtitle"])
        p.runs[0].font.size = Pt(14)
        p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    if data.get("author"):
        doc.add_paragraph(f"Author: {data['author']}")

    doc.add_paragraph("")  # spacer

    # Sections
    for section in data.get("sections", []):
        if section.get("heading"):
            h = doc.add_heading(section["heading"], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

        if section.get("content"):
            doc.add_paragraph(section["content"])

        if section.get("table"):
            tbl_data = section["table"]
            headers = tbl_data["headers"]
            rows = tbl_data["rows"]

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

            # Data rows
            for r_idx, row in enumerate(rows):
                for c_idx, val in enumerate(row):
                    table.rows[r_idx + 1].cells[c_idx].text = str(val)

            doc.add_paragraph("")  # spacer after table

        if section.get("bullet_points"):
            for point in section["bullet_points"]:
                doc.add_paragraph(point, style="List Bullet")

    doc.save(str(output_path))
    print(f"Report saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python docx-report-from-template.py report_data.json [output.docx]")
        sys.exit(1)

    data_path = Path(sys.argv[1])
    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    output = sys.argv[2] if len(sys.argv) > 2 else str(data_path.with_suffix(".docx"))
    generate_report(data, output)
